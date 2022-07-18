from functools import cmp_to_key
import sys
sys.path.append('../../project-protein-fold/RF/')
import Globals
from collections import Counter

AA_seqs = {}
Di_seqs = Globals.initialize_3Di_dict(Globals.initialize_protein_list())
tmdict = {'3':0, '5':1, '6':2, '7':3}


def categorize(AA):
    AA = AA.replace('A', 'a').replace('G', 'a').replace('V', 'a')
    AA = AA.replace('I', 'b').replace('L', 'b').replace('F', 'b').replace('P', 'b')
    AA = AA.replace('Y', 'c').replace('M', 'c').replace('T', 'c').replace('S', 'c')
    AA = AA.replace('H', 'd').replace('N', 'd').replace('Q', 'd').replace('W', 'd')
    AA = AA.replace('R', 'e').replace('K', 'e')
    AA = AA.replace('D', 'f').replace('E', 'f')
    AA = AA.replace('C', 'g')
    return AA

with open('TM_alignments/TM3_align.txt') as f:
    lines = f.readlines()
    for line in lines:
        if line[0] == '>':
            protein = line[1:].replace("\n", "")
        else:
            sequence = line.replace("\n", "")
            AA_seqs[protein] = []
            AA_seqs[protein].append(categorize(sequence))

with open('TM_alignments/TM5_align.txt') as f:
    lines = f.readlines()
    for line in lines:
        if line[0] == '>':
            protein = line[1:].replace("\n", "")
        else:
            sequence = line.replace("\n", "")
            AA_seqs[protein].append(categorize(sequence))

with open('TM_alignments/TM6_align.txt') as f:
    lines = f.readlines()
    for line in lines:
        if line[0] == '>':
            protein = line[1:].replace("\n", "")
        else:
            sequence = line.replace("\n", "")
            AA_seqs[protein].append(categorize(sequence))

with open('TM_alignments/TM7_align.txt') as f:
    lines = f.readlines()
    for line in lines:
        if line[0] == '>':
            protein = line[1:].replace("\n", "")
        else:
            sequence = line.replace("\n", "")
            AA_seqs[protein].append(categorize(sequence))

for protein in list(AA_seqs.keys()):
    j=0
    for i in range(len(AA_seqs[protein])):
        newseq = ""
        for k in range(len(AA_seqs[protein][i])):
            if AA_seqs[protein][i][k] == "-":
                newseq += "-"
            else:
                newseq += Di_seqs[protein][i][j]
                j += 1
        Di_seqs[protein][i] = newseq
        j = 0

def find_feature(tm, seq, dictionary):
    ret = []
    tmind = tmdict[tm]
    for AA in list(dictionary.keys()):
        index = 0
        while index >= 0:
            index = dictionary[AA][tmind].find(seq, index)
            if index >= 0:
                ret.append(index)
                index += 1
    if len(ret) == 0:
        return {}
    residues = []
    retcount = Counter(ret)
    temp = retcount.most_common(1)[0][1]
    for num in ret:
        if ret.count(num) == temp:
            residues.append(num)
    return set(residues)

with open('Feature_Importance/sulfur_importance.txt') as f:
    lines = f.readlines()
    i = 0
    ret = {}
    for line in lines:
        if i == 50:
            break
        i+=1
        line = line.replace('\n', "")
        if 'TM' in line:
            if line[-1] not in ret:
                ret[line[-1]] = []
            if line[0].isupper():
                for num in list(find_feature(line[-1], line[0:5], Di_seqs)):
                    ret[line[-1]].append(num)
            else:
                for num in list(find_feature(line[-1], line[0:5], AA_seqs)):
                    ret[line[-1]].append(num)
    print(ret)


#Highlighting the important regions in the TM alignments

import docx
from docx.enum.text import WD_COLOR_INDEX

doc = docx.Document()
doc.add_heading('Sulfur TM6 Important Residues', 0)

with open('TM_alignments/TM6_align.txt') as f:
    lines = f.readlines()
    for line in lines:
        paragraph = doc.add_paragraph()
        if line[0] == '>':
            paragraph.add_run(line)
        else:
            previous = 0
            for important_feature in sorted(ret['6']):
                if previous < important_feature:
                    paragraph.add_run(line[previous:important_feature])
                    paragraph.add_run(line[important_feature:important_feature + 5]).font.highlight_color = WD_COLOR_INDEX.YELLOW
                    previous = important_feature + 5
                else:
                    paragraph.add_run(line[previous:important_feature + 5]).font.highlight_color = WD_COLOR_INDEX.YELLOW
                    previous = important_feature + 5
            paragraph.add_run(line[previous:])

doc.save('TM6_highlight.docx')

